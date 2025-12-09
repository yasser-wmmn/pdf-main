import os
import re
from datetime import datetime
import fitz  # PyMuPDF
from flask import Flask, render_template, request, send_from_directory, session, abort
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from flask_httpauth import HTTPBasicAuth

# 🛠️ إعداد Tesseract OCR
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# 🛠️ تحديد مسار Poppler يدويًا
POPPLER_PATH = r"C:\Program Files\Poppler\Library\bin"

# 🔐 إعدادات الأمان
USERNAME = 'admin'           # يمكن تغييرها
PASSWORD = '123123'  # غيرها فورًا!

app = Flask(__name__)
app.secret_key = 'a_very_secure_random_string_here'  # غيرها!
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB

# إنشاء المجلدات
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# إعدادات الخط العربي
ARABIC_FONT = 'Arial'
FONT_SIZES = {
    'title': 16,
    'heading1': 14,
    'heading2': 12,
    'normal': 11,
    'small': 9
}

# محاولة استيراد المكتبات العربية
try:
    from arabic_reshaper import reshape
    from bidi.algorithm import get_display
except ImportError:
    reshape = None
    get_display = None

# تهيئة المصادقة
auth = HTTPBasicAuth()

@auth.verify_password
def verify_password(username, password):
    if username == USERNAME and password == PASSWORD:
        session['authenticated'] = True
        return username
    return None

@auth.error_handler
def unauthorized():
    return "غير مصرح بالوصول. يرجى تسجيل الدخول.", 401

def clean_arabic_text(text):
    """تنظيف النص العربي وتحسينه للعرض والنسخ"""
    if not text or not isinstance(text, str):
        return ""
    # إزالة التشكيل
    text = re.sub(r'[\u0610-\u061A\u064B-\u065F\u06D6-\u06DC\u06DF-\u06E8\u06EA-\u06ED]', '', text)
    # تصحيح الأحرف
    replacements = {
        'أ': 'ا', 'إ': 'ا', 'آ': 'ا', 'ة': 'ه', 'ى': 'ي',
        '٠': '0', '١': '1', '٢': '2', '٣': '3', '٤': '4',
        '٥': '5', '٦': '6', '٧': '7', '٨': '8', '٩': '9',
        'ـ': '', 
        'ﷲ': 'الله', 
        'هللا': 'الله', 
        'ﷺ': 'صلى الله عليه وسلم',
        # 'اال': 'الا',
        # 'اإل': 'الا',
        # 'األ': 'الا',
    }
    for wrong, correct in replacements.items():
        text = text.replace(wrong, correct)
    # تنظيف الرموز غير المرغوب فيها
    text = re.sub(r'[^\w\s\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF.,،:;؟!()\-+]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    # إعادة تشكيل النص للعرض الصحيح (مهم للعربية)
    if reshape and get_display:
        try:
            reshaped_text = reshape(text)
            return get_display(reshaped_text)
        except Exception:
            pass
    return text

def detect_text_style(text, avg_length=50):
    if not text.strip():
        return 'normal'
    if len(text) < 20 and text.isupper():
        return 'heading1'
    if text.endswith((':', '؛')) and len(text) < avg_length:
        return 'heading2'
    if len(text.split()) < 8 and any(word.isupper() for word in text.split()[:3]):
        return 'heading2'
    return 'normal'

def extract_text_blocks_with_position(page):
    blocks = []
    for block in page.get_text("blocks"):
        x0, y0, x1, y1, text, block_no, block_type = block
        if text.strip():
            blocks.append({
                'text': text.strip(),
                'x0': x0, 'y0': y0, 'x1': x1, 'y1': y1
            })
    return blocks

def sort_blocks_reading_order(blocks, y_threshold=10):
    if not blocks:
        return []
    sorted_blocks = sorted(blocks, key=lambda b: (b['y0'], -b['x0']))
    lines = []
    for block in sorted_blocks:
        placed = False
        for line in lines:
            if abs(block['y0'] - line[0]['y0']) < y_threshold:
                line.append(block)
                placed = True
                break
        if not placed:
            lines.append([block])
    lines.sort(key=lambda l: l[0]['y0'])
    result_blocks = []
    for line in lines:
        sorted_line = sorted(line, key=lambda b: -b['x0'])
        result_blocks.extend(sorted_line)
    return result_blocks

def extract_text_from_pdf(pdf_path):
    doc_data = {'title': '', 'author': '', 'pages': []}
    custom_config = r'--oem 3 --psm 6 -l ara+eng'
    try:
        with fitz.open(pdf_path) as pdf:
            doc_data['title'] = pdf.metadata.get('title', '') or os.path.basename(pdf_path)
            doc_data['author'] = pdf.metadata.get('author', '')
            for page_num in range(len(pdf)):
                page = pdf[page_num]
                page_content = []
                blocks = extract_text_blocks_with_position(page)
                if blocks:
                    ordered_blocks = sort_blocks_reading_order(blocks)
                    for block in ordered_blocks:
                        text = clean_arabic_text(block['text'])
                        if not text:
                            continue
                        style = detect_text_style(text)
                        page_content.append({'text': text, 'style': style})
                else:
                    try:
                        images = convert_from_path(
                            pdf_path,
                            first_page=page_num + 1,
                            last_page=page_num + 1,
                            poppler_path=POPPLER_PATH
                        )
                        for img in images:
                            ocr_text = pytesseract.image_to_string(img, config=custom_config)
                            for line in ocr_text.split('\n'):
                                line = line.strip()
                                if not line:
                                    continue
                                line = clean_arabic_text(line)
                                style = detect_text_style(line)
                                page_content.append({'text': line, 'style': style})
                    except Exception as e:
                        print(f"OCR failed on page {page_num + 1}: {e}")
                        page_content.append({'text': '[فشل في قراءة الصفحة]', 'style': 'normal'})
                doc_data['pages'].append({
                    'number': page_num + 1,
                    'content': page_content
                })
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {e}")
        abort(500)
    return doc_data

def extract_text_from_image(image_path):
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image, lang='ara+eng')
        return text
    except Exception as e:
        print(f"OCR Error: {e}")
        return "فشل في قراءة الصورة."

def add_to_document(doc, text, style='normal'):
    """إضافة نص إلى المستند مع التنسيق المناسب (يمين لليسار، خط عربي)"""
    if not text.strip():
        return

    # 1. إنشاء فقرة جديدة
    paragraph = doc.add_paragraph()
    
    # 2. تطبيق اتجاه النص: من اليمين لليسار
    paragraph_format = paragraph.paragraph_format
    paragraph_format.right_to_left = True  # ⭐ هذا مهم جدًا
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 3. إضافة النص كـ run
    run = paragraph.add_run(text)

    # 4. ضبط الخط العربي
    run.font.name = ARABIC_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), ARABIC_FONT)
    run.font.size = Pt(FONT_SIZES[style])

    # 5. تنسيق حسب النوع
    if style == 'title':
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 139)
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format.right_to_left = False  # العنوان في المنتصف، ليس RTL
    elif style == 'heading1':
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif style == 'heading2':
        run.bold = True
        run.font.color.rgb = RGBColor(34, 139, 34)
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)

        
def create_google_docs_like_document(data, output_path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    title = data.get('title', f"مستند مفرغ - {datetime.now().strftime('%Y-%m-%d')}")
    add_to_document(doc, title, 'title')
    if data.get('author'):
        add_to_document(doc, f"المؤلف: {data['author']}", 'small')
    add_to_document(doc, f"تاريخ التفريغ: {datetime.now().strftime('%Y-%m-%d %H:%M')}", 'small')
    add_to_document(doc, "—" * 50, 'small')
    for page in data['pages']:
        add_to_document(doc, f"الصفحة {page['number']}", 'heading1')
        for content in page['content']:
            add_to_document(doc, content['text'], content['style'])
        if page['number'] < len(data['pages']):
            add_to_document(doc, "—" * 50, 'small')
    doc.save(output_path)

def allowed_file(filename):
    allowed = {'pdf', 'jpg', 'jpeg', 'png', 'bmp', 'tiff', 'tif'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed

@app.route('/', methods=['GET', 'POST'])
@auth.login_required
def index():
    if request.method == 'POST':
        if 'pdf_file' not in request.files:
            return render_template('index.html', error="لم يتم اختيار ملف")
        file = request.files['pdf_file']
        if file.filename == '':
            return render_template('index.html', error="لم يتم اختيار ملف")
        if not allowed_file(file.filename):
            return render_template('index.html', error="الملف غير مدعوم.")
        try:
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            ext = filename.rsplit('.', 1)[1].lower()
            if ext == 'pdf':
                data = extract_text_from_pdf(file_path)
            else:
                raw_text = extract_text_from_image(file_path)
                lines = [clean_arabic_text(line.strip()) for line in raw_text.split('\n') if line.strip()]
                content = [{'text': line, 'style': detect_text_style(line)} for line in lines]
                data = {
                    'title': filename,
                    'author': '',
                    'pages': [{'number': 1, 'content': content}]
                }

            output_filename = f"مستند_{os.path.splitext(filename)[0]}.docx"
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            create_google_docs_like_document(data, output_path)

            # ✅ تنظيف الملفات المؤقتة بعد المعالجة
            if os.path.exists(file_path):
                os.remove(file_path)

            download_links = [f"/download/{output_filename}"]
            stats = {
                'pages': len(data['pages']),
                'words': sum(len(re.findall(r'\w+', c['text'])) for p in data['pages'] for c in p['content'])
            }
            return render_template('index.html',
                                 download_links=download_links,
                                 document_title=data['title'],
                                 stats=stats)
        except Exception as e:
            return render_template('index.html', error=f"حدث خطأ: {str(e)}")
    return render_template('index.html')

@app.route('/download/<filename>')
@auth.login_required
def download(filename):
    file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    if not os.path.exists(file_path):
        abort(404)
    return send_from_directory(app.config['OUTPUT_FOLDER'], filename, as_attachment=True)

# ❌ تعطيل المعالجة الدُفعية تلقائيًا (لأنها قد تكون خطرة)
# batch_process_pdfs()

if __name__ == '__main__':
    # ⚠️ مهم: لا تُشغّل debug=True أبدًا في بيئة حقيقية
    app.run(debug=False, host='0.0.0.0', port=5000)